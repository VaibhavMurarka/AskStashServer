# dependencies.py

from dotenv import load_dotenv
load_dotenv() # This loads the variables from .env

import hashlib
import jwt
import os
from datetime import datetime, timedelta
import google.generativeai as genai
from supabase import create_client, Client
import io
from PIL import Image
import PyPDF2
from docx import Document
import filetype # Using filetype instead of magic

# Supabase configuration
SUPABASE_URL = os.environ.get('SUPABASE_URL')
SUPABASE_KEY = os.environ.get('SUPABASE_KEY')
if not SUPABASE_URL or not SUPABASE_KEY:
    raise RuntimeError("Supabase URL and Key must be set in environment variables.")
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

# JWT Configuration
SECRET_KEY = os.environ.get('SECRET_KEY')
if not SECRET_KEY:
    raise RuntimeError("SECRET_KEY must be set in environment variables.")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 30

# Initialize Gemini AI with your API key
GEMINI_API_KEY = os.environ.get('GEMINI_API_KEY')
if not GEMINI_API_KEY:
    raise RuntimeError("GEMINI_API_KEY must be set in environment variables.")
genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel('gemini-1.5-flash')

# --- Utility Functions ---

def hash_password(password: str) -> str:
    return hashlib.sha256(password.encode()).hexdigest()

def verify_password(plain_password: str, hashed_password: str) -> bool:
    return hash_password(plain_password) == hashed_password

def create_access_token(data: dict, expires_delta: timedelta = None):
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=15)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """Text extraction function using the 'filetype' library."""
    try:
        kind = filetype.guess(file_content)
        file_type = kind.mime if kind else 'application/octet-stream'
        file_extension = filename.lower().split('.')[-1] if '.' in filename else ''

        if file_type == 'application/pdf' or file_extension == 'pdf':
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                text_parts = []
                for i, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_parts.append(f"\n--- Page {i + 1} ---\n{page_text}\n")
                return "".join(text_parts).strip()
            except Exception as e:
                return f"[Error reading PDF: {str(e)}]"

        elif (file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or
              file_extension == 'docx'):
            try:
                doc = Document(io.BytesIO(file_content))
                full_text = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                full_text.append(cell.text)
                return '\n'.join(full_text).strip()
            except Exception as e:
                return f"[Error reading Word document: {str(e)}]"

        elif file_extension == 'doc':
            return f"[Legacy Word document (.doc) detected. Please convert to .docx format for better text extraction. Filename: {filename}]"

        elif (file_type.startswith('text/') or
              file_extension in ['txt', 'md', 'csv', 'json', 'xml', 'html', 'htm', 'rtf']):
            try:
                return file_content.decode('utf-8')
            except UnicodeDecodeError:
                return file_content.decode('latin-1', errors='ignore')
            except Exception as e:
                return f"[Error reading text file: {str(e)}]"

        elif (file_type.startswith('image/') or
              file_extension in ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'webp', 'tiff']):
            try:
                vision_model = genai.GenerativeModel('gemini-1.5-flash')
                image = Image.open(io.BytesIO(file_content))
                prompt = "Extract all text from this image. If no text is present, briefly describe the image content."
                response = vision_model.generate_content([prompt, image])
                return f"[Image content from {filename}]\n{response.text}"
            except Exception as e:
                return f"[Image content from {filename} - Vision processing error: {str(e)}]"

        else:
            return f"[Unsupported file type: {file_type} (extension: .{file_extension}).]"

    except Exception as e:
        return f"[Error extracting text from {filename}: {str(e)}]"

def generate_response(message: str, context: str = "") -> str:
    """AI response generation function."""
    try:
        if context and context.strip():
            prompt = f"""You are a helpful AI assistant with access to the user's uploaded documents.

CONTEXT FROM USER'S DOCUMENTS:
{context}

USER QUESTION: {message}

INSTRUCTIONS:
1. Answer the user's question primarily based on the provided document context.
2. If the documents contain relevant information, cite specific parts and be detailed.
3. If the documents don't contain relevant information, clearly state this and provide general knowledge.

RESPONSE:"""
        else:
            prompt = f"""You are a helpful AI assistant. The user has not selected any documents for context.

USER QUESTION: {message}

Please provide a helpful, informative response based on your general knowledge.

RESPONSE:"""
        response = model.generate_content(
            prompt,
            generation_config={
                'temperature': 0.7,
                'top_p': 0.8,
                'top_k': 40,
                'max_output_tokens': 2048,
            }
        )
        return response.text
    except Exception as e:
        return f"I apologize, but I encountered an error while processing your request: {str(e)}"
    
    